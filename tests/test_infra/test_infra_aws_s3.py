from docx import Document

from app.infra.repository.s3 import S3Repository
from tests.conftest import create_s3_test_data, create_s3_test_data_for_docx


def test_s3_file_exists() -> None:
    bucket_name = "test-bucket"
    test_files = {
        "test-folder/sample.txt": "This is a test file.",
        "test-folder/another.txt": "Another test file.",
    }

    s3, bucket_name = create_s3_test_data(bucket_name, test_files)

    response = s3.list_objects_v2(Bucket=bucket_name)
    keys = [obj["Key"] for obj in response.get("Contents", [])]

    for expected_key in test_files.keys():
        assert expected_key in keys, f"ファイル {expected_key} が S3 に存在しません"

    # クリーンアップ
    for key in test_files.keys():
        s3.delete_object(Bucket=bucket_name, Key=key)
    s3.delete_bucket(Bucket=bucket_name)


def test_get_docx_file() -> None:
    bucket_name = "test-docx-bucket"
    key = "documents/test.docx"

    # テストデータをアップロード
    doc = Document()
    doc.add_paragraph("Hello, this is a test DOCX file.")
    s3, _ = create_s3_test_data_for_docx(bucket_name, {key: doc})

    # テスト対象のリポジトリ
    s3_repo = S3Repository()

    # テスト実行
    loaded_doc = s3_repo.get_docx_file(
        key="documents/test.docx", bucket_name=bucket_name
    )
    assert loaded_doc.paragraphs[0].text == "Hello, this is a test DOCX file."

    # クリーンアップ
    s3.delete_object(Bucket=bucket_name, Key=key)
    s3.delete_bucket(Bucket=bucket_name)
